import os
import markdown
from bs4 import BeautifulSoup


def read_file_with_auto_encoding(file_path):
    encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read(), encoding
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("无法以任何已知编码读取文件")


# 获取当前目录
current_dir = os.getcwd()

# 遍历当前目录下的所有文件
for filename in os.listdir(current_dir):
    if filename.endswith('.md'):
        file_path = os.path.join(current_dir, filename)
        try:
            # 尝试读取Markdown文件内容
            md_content, detected_encoding = read_file_with_auto_encoding(file_path)

            # 将Markdown转换为HTML，启用额外的扩展
            html_content = markdown.markdown(md_content,
                                             extensions=['fenced_code', 'codehilite', 'nl2br', 'sane_lists'])

            # 创建一个基本的HTML结构
            soup = BeautifulSoup('<html><head></head><body></body></html>', 'html.parser')

            # 设置HTML的编码
            meta_charset = soup.new_tag('meta')
            meta_charset.attrs['charset'] = 'utf-8'
            soup.head.append(meta_charset)

            # 将转换后的内容添加到<body>标签中
            soup.body.append(BeautifulSoup(html_content, 'html.parser'))

            # 添加样式
            style = soup.new_tag('style')
            style.string = '''
                body { font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }
                h1, h2, h3 { color: #333; }
                pre { background-color: #f4f4f4; border: 1px solid #ddd; border-radius: 4px; padding: 10px; overflow-x: auto; }
                code { font-family: 'Courier New', Courier, monospace; }
                .codehilite { background: #f8f8f8; }
                ul, ol { padding-left: 30px; }
                li { margin-bottom: 10px; }
            '''
            soup.head.append(style)

            # 生成HTML文件名
            html_filename = os.path.splitext(filename)[0] + '.html'

            # 将HTML内容写入文件
            with open(html_filename, 'w', encoding='utf-8') as html_file:
                html_file.write(str(soup))

            print(f'已将 {filename} (编码: {detected_encoding}) 转换为 {html_filename}')
        except UnicodeDecodeError:
            print(f'无法读取文件 {filename}，跳过')

print('转换完成！')
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from htmldocx import HtmlToDocx

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), "DDDDDD")
    p._element.get_or_add_pPr().append(shading_elm)

def html_to_docx(html_file, docx_file):
    # 读取HTML文件
    with open(html_file, 'r', encoding='utf-8') as file:
        html_content = file.read()

    # 创建Word文档对象
    doc = Document()

    # 解析HTML内容
    soup = BeautifulSoup(html_content, 'html.parser')

    # 获取HTML文件所在目录
    html_dir = os.path.dirname(html_file)

    # 处理HTML内容
    for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'pre', 'ul', 'ol']):
        if elem.name == 'pre':
            code = elem.find('code')
            if code:
                add_code_block(doc, code.text)
        elif elem.name in ['ul', 'ol']:
            for li in elem.find_all('li'):
                p = doc.add_paragraph(li.text, style='List Bullet' if elem.name == 'ul' else 'List Number')
        else:
            if elem.name.startswith('h'):
                level = int(elem.name[1])
                p = doc.add_heading(level=level)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for content in elem.contents:
                if isinstance(content, str):
                    p.add_run(content)
                elif content.name == 'img':
                    img_src = content.get('src')
                    img_path = os.path.join(html_dir, img_src)
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6))
                    else:
                        p.add_run(f"[图片未找到: {img_src}]")
                elif content.name == 'code':
                    run = p.add_run(content.text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                else:
                    p.add_run(content.text)

    # 设置中文字体
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存DOCX文件
    doc.save(docx_file)

# 获取当前目录下所有的HTML文件
html_files = [f for f in os.listdir('.') if f.endswith('.html')]

# 转换每个HTML文件为DOCX
for html_file in html_files:
    docx_file = os.path.splitext(html_file)[0] + '.docx'
    print(f"正在转换 {html_file} 为 {docx_file}...")
    html_to_docx(html_file, docx_file)

print("转换完成！")
