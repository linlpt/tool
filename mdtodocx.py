import os
import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from bs4 import BeautifulSoup
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def convert_md_to_docx(md_file, docx_file):
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 将Markdown转换为HTML
    html = markdown.markdown(md_content, extensions=['fenced_code', 'tables', 'nl2br'])

    # 创建Word文档
    doc = Document()

    # 解析HTML
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'pre', 'code', 'img']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            paragraph = doc.add_heading(element.text, level=level)
        elif element.name == 'p':
            # 检查段落是否只包含图片
            if element.find('img') and len(element.contents) == 1:
                img = element.find('img')
                add_image_to_doc(doc, img['src'], md_file)
            else:
                paragraph = doc.add_paragraph(element.get_text())
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                # 检查是否为编号列表项
                match = re.match(r'^(\d+\.)\s*(.+)$', li.text.strip())
                if match and element.name == 'ol':
                    number, content = match.groups()
                    paragraph = doc.add_paragraph(style='List Number')
                    paragraph.add_run(f"{number} {content}")
                else:
                    paragraph = doc.add_paragraph(li.text.strip(), style='List Bullet' if element.name == 'ul' else 'List Number')
        elif element.name == 'pre':
            code_text = element.text.strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Inches(0.1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.right_indent = Inches(0.5)
            paragraph.style = 'No Spacing'
            # 添加边框
            set_border(paragraph)
        elif element.name == 'img':
            img_src = element['src']
            add_image_to_doc(doc, img_src, md_file)

    # 保存Word文档
    doc.save(docx_file)

def add_image_to_doc(doc, img_src, md_file):
    if img_src.startswith('http'):
        doc.add_picture(img_src, width=Inches(6))
    else:
        img_path = os.path.join(os.path.dirname(md_file), img_src)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))  # 设置图片宽度为6英寸
        else:
            print(f"警告：找不到图片文件 {img_path}")

def set_border(paragraph):
    pBdr = OxmlElement('w:pBdr')
    borders = [
        ('w:top', {'w:val': 'single', 'w:sz': '4', 'w:space': '1', 'w:color': 'auto'}),
        ('w:left', {'w:val': 'single', 'w:sz': '4', 'w:space': '1', 'w:color': 'auto'}),
        ('w:bottom', {'w:val': 'single', 'w:sz': '4', 'w:space': '1', 'w:color': 'auto'}),
        ('w:right', {'w:val': 'single', 'w:sz': '4', 'w:space': '1', 'w:color': 'auto'}),
    ]
    for name, attrs in borders:
        element = OxmlElement(name)
        for k, v in attrs.items():
            element.set(qn(k), v)
        pBdr.append(element)
    paragraph._element.get_or_add_pPr().append(pBdr)

# 获取当前目录下所有的md文件
md_files = [f for f in os.listdir('.') if f.endswith('.md')]

# 转换所有md文件为docx
for md_file in md_files:
    docx_file = md_file[:-3] + '.docx'
    convert_md_to_docx(md_file, docx_file)
    print(f'已将 {md_file} 转换为 {docx_file}')